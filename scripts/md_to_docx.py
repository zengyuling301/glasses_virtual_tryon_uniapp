#!/usr/bin/env python3
"""
Markdown → DOCX.

1) If ``pandoc`` is on PATH, uses pandoc (best layout; install: brew install pandoc).
2) Otherwise builds the document with python-docx + BeautifulSoup (headings, tables,
   lists, horizontal rules, CJK-friendly body font).
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any, Iterator

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _try_pandoc(src: Path, out: Path) -> bool:
    exe = shutil.which("pandoc")
    if not exe:
        return False
    try:
        subprocess.run(
            [exe, str(src), "-f", "markdown", "-t", "docx", "-o", str(out)],
            check=True,
            capture_output=True,
            text=True,
        )
        return True
    except (subprocess.CalledProcessError, OSError):
        return False


def _set_run_east_asia(run: Any, western: str = "Calibri", east_asia: str = "Microsoft YaHei") -> None:
    run.font.name = western
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), western)
        rfonts.set(qn("w:hAnsi"), western)
        rfonts.set(qn("w:eastAsia"), east_asia)
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:eastAsia"), east_asia)
        rfonts.set(qn("w:ascii"), western)
        rfonts.set(qn("w:hAnsi"), western)


def _iter_inline_chunks(
    node: Tag,
    *,
    bold: bool = False,
    italic: bool = False,
    mono: bool = False,
) -> Iterator[tuple[str, bool, bool, bool]]:
    for child in node.children:
        if isinstance(child, NavigableString):
            s = str(child)
            if s:
                yield s, bold, italic, mono
            continue
        if not isinstance(child, Tag):
            continue
        name = child.name.lower()
        if name in ("strong", "b"):
            yield from _iter_inline_chunks(child, bold=True, italic=italic, mono=mono)
        elif name in ("em", "i"):
            yield from _iter_inline_chunks(child, bold=bold, italic=True, mono=mono)
        elif name == "code":
            yield from _iter_inline_chunks(child, bold=bold, italic=italic, mono=True)
        elif name == "br":
            yield "\n", bold, italic, mono
        elif name == "a":
            yield from _iter_inline_chunks(child, bold=bold, italic=italic, mono=mono)
        else:
            yield from _iter_inline_chunks(child, bold=bold, italic=italic, mono=mono)


def _add_inline_runs(paragraph: Any, node: Tag) -> None:
    for text, bold, italic, mono in _iter_inline_chunks(node):
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if mono:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            rpr = run._element.get_or_add_rPr()
            rf = rpr.rFonts
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rf.set(qn("w:ascii"), "Consolas")
                rf.set(qn("w:hAnsi"), "Consolas")
                rf.set(qn("w:eastAsia"), "SimSun")
                rpr.append(rf)
            else:
                rf.set(qn("w:ascii"), "Consolas")
                rf.set(qn("w:hAnsi"), "Consolas")
                rf.set(qn("w:eastAsia"), "SimSun")
        else:
            _set_run_east_asia(run)


def _add_horizontal_rule(document: Document) -> None:
    p = document.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="6" w:space="2" w:color="BBBBBB"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def _add_table(document: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr")
    if not rows:
        return
    col_counts: list[int] = []
    for tr in rows:
        cells = tr.find_all(("th", "td"), recursive=False)
        col_counts.append(len(cells))
    ncols = max(col_counts) if col_counts else 1
    nrows = len(rows)
    tbl = document.add_table(rows=nrows, cols=ncols)
    tbl.style = "Table Grid"
    for ri, tr in enumerate(rows):
        cells = tr.find_all(("th", "td"), recursive=False)
        for ci, cell in enumerate(cells):
            if ci >= ncols:
                break
            tc = tbl.rows[ri].cells[ci]
            tc.text = ""
            p = tc.paragraphs[0]
            if cell.name == "th":
                _add_inline_runs(p, cell)
                for r in p.runs:
                    r.bold = True
            else:
                _add_inline_runs(p, cell)
            for r in p.runs:
                _set_run_east_asia(r)
                r.font.size = Pt(10)


def _clear_paragraph_runs(paragraph: Any) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith(("r", "hyperlink")):
            p.remove(child)


def _walk_block(document: Document, node: Tag) -> None:
    name = node.name.lower()
    if name == "h1":
        p = document.add_heading("", level=1)
        _clear_paragraph_runs(p)
        _add_inline_runs(p, node)
        for r in p.runs:
            _set_run_east_asia(r, western="Calibri Light", east_asia="Microsoft YaHei")
    elif name == "h2":
        p = document.add_heading("", level=2)
        _clear_paragraph_runs(p)
        _add_inline_runs(p, node)
        for r in p.runs:
            _set_run_east_asia(r)
    elif name == "h3":
        p = document.add_heading("", level=3)
        _clear_paragraph_runs(p)
        _add_inline_runs(p, node)
        for r in p.runs:
            _set_run_east_asia(r)
    elif name == "p":
        p = document.add_paragraph()
        _add_inline_runs(p, node)
        for r in p.runs:
            _set_run_east_asia(r)
            r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
    elif name == "ul":
        for li in node.find_all("li", recursive=False):
            p = document.add_paragraph(style="List Bullet")
            _clear_paragraph_runs(p)
            _add_inline_runs(p, li)
            for r in p.runs:
                _set_run_east_asia(r)
                r.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(2)
    elif name == "ol":
        for li in node.find_all("li", recursive=False):
            p = document.add_paragraph(style="List Number")
            _clear_paragraph_runs(p)
            _add_inline_runs(p, li)
            for r in p.runs:
                _set_run_east_asia(r)
                r.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(2)
    elif name == "table":
        _add_table(document, node)
        document.add_paragraph()
    elif name == "hr":
        _add_horizontal_rule(document)
    elif name == "pre":
        txt = node.get_text()
        p = document.add_paragraph(txt)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(9)
    elif name in ("div", "blockquote"):
        for child in node.children:
            if isinstance(child, Tag):
                _walk_block(document, child)


def _build_docx_with_docx(src: Path, out: Path, title: str) -> None:
    text = src.read_text(encoding="utf-8")
    body_html = markdown.markdown(text, extensions=["tables", "fenced_code", "nl2br"])
    soup = BeautifulSoup(f"<div id='mdroot'>{body_html}</div>", "html.parser")
    root = soup.find("div", id="mdroot")
    if root is None:
        raise RuntimeError("Failed to parse markdown HTML wrapper.")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    cp = doc.core_properties
    cp.title = title
    cp.subject = title

    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    if normal.font is not None:
        normal.font.name = "Calibri"

    for child in root.children:
        if isinstance(child, Tag):
            _walk_block(doc, child)

    doc.save(str(out))


def main() -> None:
    p = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("input", type=Path, help="Input .md path")
    p.add_argument("-o", "--output", type=Path, default=None, help="Output .docx (default: same stem as input)")
    p.add_argument("--title", type=str, default="", help="Core document title (python-docx path only)")
    p.add_argument("--no-pandoc", action="store_true", help="Skip pandoc even if installed")
    args = p.parse_args()
    src: Path = args.input
    if not src.is_file():
        print(f"Not found: {src}", file=sys.stderr)
        sys.exit(1)
    out: Path = args.output if args.output else src.with_suffix(".docx")
    title = args.title or src.stem

    if not args.no_pandoc and _try_pandoc(src, out):
        print(out.resolve(), "(pandoc)")
        return

    try:
        _build_docx_with_docx(src, out, title)
    except ImportError as e:
        print(
            "Missing dependency. Install: pip install python-docx beautifulsoup4 markdown",
            file=sys.stderr,
        )
        raise e
    print(out.resolve(), "(python-docx)")


if __name__ == "__main__":
    main()
